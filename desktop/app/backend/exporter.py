"""Unified exporter for the three artifact kinds — raw transcript, summary
(markdown) and analysis (the 11-feature JSON) — into a common set of formats:
``txt, md, json, html, pdf, docx``.

Design goals (owner requirements):
- ONE rendering pass per artifact builds a tiny block model
  (h/p/ul/ol/kv/rule/footer); every format writer consumes the same blocks, so
  structure is identical across formats and NO field is dropped.
- Every exported file ends with an app-name + version footer.
- Naming is version-aware via ``versioned_filename`` so the exported version
  always matches the one shown in the UI (no "v3 in UI, v1 on disk").
- Analysis renders ALL 11 features with the real schema field names (the bug in
  the legacy exporter dropped 9 of them and misread the rest).

JSON is special-cased (a structured dump, not block-rendered) so integrations
get the full, loss-free structure.
"""
from __future__ import annotations

import html as _html
import json
import re
from datetime import datetime
from pathlib import Path

from ..core.history import versioned_filename
from ..version import APP_VERSION

APP_NAME = "Meeting Summarizer"
FORMATS = ("txt", "md", "json", "html", "pdf", "docx")
KINDS = ("raw", "summary", "analysis")


def app_version() -> str:
    """Read the version from package.json (single source of truth)."""
    return APP_VERSION


def footer_text() -> str:
    return f"{APP_NAME} v{app_version()}"


def default_export_path(out_dir, stem: str, kind: str, version: int, fmt: str) -> Path:
    """`<stem>_<kind>[ _vN ].<fmt>` — raw is never versioned (version forced 1,
    so it is always `<stem>_raw.<fmt>`)."""
    if kind == "raw":
        version = 1
    return Path(out_dir) / versioned_filename(stem, kind, version, "." + fmt)


# -- analysis field labels (ru default; en available) ----------------------
LBL = {
    "ru": {
        "characteristics": "Характеристики встречи", "keyTopics": "Ключевые темы",
        "frequentWords": "Частые слова", "actionItems": "Задачи и действия",
        "assignee": "Ответственный", "priority": "Приоритет", "deadline": "Срок",
        "sentiment": "Тональность и настроение", "overall": "Общая тональность",
        "engagement": "Вовлечённость", "hasConflict": "Конфликт", "emotions": "Эмоции",
        "description": "Описание", "interruptionIndex": "Индекс прерываний",
        "emotionalBalance": "Баланс эмоций", "empathyIndex": "Индекс эмпатии",
        "speechSpeedVariability": "Вариативность темпа речи",
        "questionsToAnswersRatio": "Вопросы/ответы",
        "dominanceDistribution": "Доминирование", "category": "Категория встречи",
        "tags": "Теги", "risks": "Риски и блокеры", "severity": "Серьёзность",
        "impact": "Влияние", "status": "Статус", "quotes": "Ключевые цитаты",
        "speaker": "Спикер", "context": "Контекст", "technologies": "Технологии и системы",
        "questions": "Открытые вопросы", "owner": "Ответ за", "recommendations": "Рекомендации",
        "followupQuestions": "Вопросы к следующей встрече", "formalProtocol": "Формальный протокол",
        "protocolNumber": "Номер протокола", "date": "Дата", "time": "Время",
        "location": "Место", "participants": "Участники", "chairman": "Председатель",
        "secretary": "Секретарь", "agenda": "Повестка", "decisions": "Решения",
        "decision": "Решение", "votingResult": "Голосование", "nextMeeting": "Следующая встреча",
        "protocolText": "Текст протокола", "yes": "да", "no": "нет", "name": "Название",
        "task": "Задача", "question": "Вопрос", "recommendation": "Рекомендация",
        "duration": "Длительность", "wordCount": "Количество слов",
        "analysisReport": "Анализ встречи",
    },
    "en": {
        "characteristics": "Meeting Characteristics", "keyTopics": "Key Topics",
        "frequentWords": "Frequent Words", "actionItems": "Action Items",
        "assignee": "Assignee", "priority": "Priority", "deadline": "Deadline",
        "sentiment": "Sentiment & Tone", "overall": "Overall sentiment",
        "engagement": "Engagement", "hasConflict": "Conflict", "emotions": "Emotions",
        "description": "Description", "interruptionIndex": "Interruption index",
        "emotionalBalance": "Emotional balance", "empathyIndex": "Empathy index",
        "speechSpeedVariability": "Speech speed variability",
        "questionsToAnswersRatio": "Questions/answers",
        "dominanceDistribution": "Dominance", "category": "Meeting Category",
        "tags": "Tags", "risks": "Risks & Blockers", "severity": "Severity",
        "impact": "Impact", "status": "Status", "quotes": "Key Quotes",
        "speaker": "Speaker", "context": "Context", "technologies": "Technologies & Systems",
        "questions": "Open Questions", "owner": "Owner", "recommendations": "Recommendations",
        "followupQuestions": "Follow-up Questions", "formalProtocol": "Formal Protocol",
        "protocolNumber": "Protocol number", "date": "Date", "time": "Time",
        "location": "Location", "participants": "Participants", "chairman": "Chairman",
        "secretary": "Secretary", "agenda": "Agenda", "decisions": "Decisions",
        "decision": "Decision", "votingResult": "Voting", "nextMeeting": "Next meeting",
        "protocolText": "Protocol text", "yes": "yes", "no": "no", "name": "Name",
        "task": "Task", "question": "Question", "recommendation": "Recommendation",
        "duration": "Duration", "wordCount": "Word count",
        "analysisReport": "Meeting Analysis",
    },
}


# -- block builders --------------------------------------------------------
# Block model (one pass, shared by every writer):
#   ("h", level, text) | ("p", text) | ("ul", [str]) | ("ol", [str])
#   ("kv", [(key, value), ...]) | ("rule",) | ("footer", text)

_NUM_RE = re.compile(r"^\d+[.)]\s+")

_TITLES = {
    "ru": {"raw": "Транскрипция", "summary": "Саммари встречи", "analysis": "Анализ встречи"},
    "en": {"raw": "Transcript", "summary": "Meeting Summary", "analysis": "Meeting Analysis"},
}
_HEAD = {
    "ru": {"video": "Видео", "date": "Дата экспорта", "duration": "Длительность",
           "version": "Версия"},
    "en": {"video": "Video", "date": "Export date", "duration": "Duration",
           "version": "Version"},
}


def markdown_to_blocks(text: str) -> list:
    blocks: list = []
    para: list = []
    lst: list = []
    lst_kind = None

    def flush_para():
        if para:
            blocks.append(("p", "\n".join(para)))
            para.clear()

    def flush_list():
        nonlocal lst, lst_kind
        if lst:
            blocks.append((lst_kind, lst))
            lst = []
            lst_kind = None

    for raw_line in (text or "").split("\n"):
        stripped = raw_line.strip()
        if not stripped:
            flush_para(); flush_list(); continue
        if stripped.startswith("#"):
            flush_para(); flush_list()
            n = len(stripped) - len(stripped.lstrip("#"))
            blocks.append(("h", min(n, 3), stripped[n:].strip()))
            continue
        if stripped.startswith("- ") or stripped.startswith("* "):
            flush_para()
            if lst_kind not in (None, "ul"):
                flush_list()
            lst_kind = "ul"; lst.append(stripped[2:].strip())
            continue
        m = _NUM_RE.match(stripped)
        if m:
            flush_para()
            if lst_kind not in (None, "ol"):
                flush_list()
            lst_kind = "ol"; lst.append(stripped[m.end():].strip())
            continue
        flush_list()
        para.append(stripped)
    flush_para(); flush_list()
    return blocks


def build_raw_blocks(text: str) -> list:
    blocks = [("p", p.strip()) for p in re.split(r"\n\s*\n", (text or "").strip()) if p.strip()]
    return blocks or [("p", "")]


# Enum VALUES the model emits in English (positive/neutral/…, high/medium/low) —
# translated for display so the RU report never shows raw "neutral"/"high".
_SENTIMENT_VAL = {
    "ru": {"positive": "Позитивная", "neutral": "Нейтральная", "negative": "Негативная"},
    "en": {"positive": "Positive", "neutral": "Neutral", "negative": "Negative"},
}
_LEVEL_VAL = {
    "ru": {"high": "Высокая", "medium": "Средняя", "low": "Низкая"},
    "en": {"high": "High", "medium": "Medium", "low": "Low"},
}


def _sentiment_val(value, lang: str):
    return _SENTIMENT_VAL.get(lang, _SENTIMENT_VAL["en"]).get(str(value).lower(), value)


def _level_val(value, lang: str):
    return _LEVEL_VAL.get(lang, _LEVEL_VAL["en"]).get(str(value).lower(), value)


def build_analysis_blocks(a: dict, lang: str = "ru", meta: dict = None) -> list:
    L = (lambda k: LBL.get(lang, LBL["en"]).get(k, k))
    a = a or {}
    meta = meta or {}
    blocks: list = []

    def add_items(key, title_key, head_keys, head_label_fn, lead_key=None, numbered=True):
        items = a.get(key) or []
        if not items:
            return
        blocks.append(("h", 2, L(title_key)))
        for i, it in enumerate(items, 1):
            if not isinstance(it, dict):
                blocks.append(("p", str(it)))
                continue
            lead = it.get(lead_key, "") if lead_key else ""
            # The item's text is user/model data.  Never rstrip punctuation from
            # it merely to format the numeric prefix: that made TXT/MD/DOCX lose
            # final periods while HTML/JSON kept them.
            head = f"{i}. {lead}" if (numbered and lead) else (lead or f"{i}.")
            blocks.append(("h", 3, head))
            kv = [(head_label_fn(k), it[k]) for k in head_keys if it.get(k) not in (None, "")]
            if kv:
                blocks.append(("kv", kv))

    # characteristics (meeting meta + keyTopics, frequentWords)
    ch = a.get("characteristics") or {}
    kt, fw = ch.get("keyTopics") or [], ch.get("frequentWords") or []
    char_kv = []
    if meta.get("duration"):
        char_kv.append((L("duration"), meta["duration"]))
    if meta.get("participants"):
        char_kv.append((L("participants"), meta["participants"]))
    if meta.get("wordCount") is not None:
        char_kv.append((L("wordCount"), f"{int(meta['wordCount']):,}".replace(",", " ")))
    if char_kv or kt or fw:
        blocks.append(("h", 2, L("characteristics")))
        if char_kv:
            blocks.append(("kv", char_kv))
        if kt:
            blocks.append(("h", 3, L("keyTopics")))
            blocks.append(("ul", [str(x) for x in kt]))
        if fw:
            blocks.append(("h", 3, L("frequentWords")))
            fwi = [f"{x[0]} — {x[1]}" if isinstance(x, (list, tuple)) and len(x) >= 2
                   else str(x) for x in fw]
            blocks.append(("ul", fwi))

    add_items("actionItems", "actionItems", ("assignee", "priority", "deadline"), L, "task")

    s = a.get("sentiment")
    if isinstance(s, dict) and s:
        blocks.append(("h", 2, L("sentiment")))
        kv = []
        if s.get("overall") not in (None, ""):
            kv.append((L("overall"), _sentiment_val(s["overall"], lang)))
        if s.get("engagement") not in (None, ""):
            kv.append((L("engagement"), _level_val(s["engagement"], lang)))
        if "hasConflict" in s:
            kv.append((L("hasConflict"), L("yes") if s.get("hasConflict") else L("no")))
        if s.get("emotions"):
            kv.append((L("emotions"), ", ".join(map(str, s["emotions"]))))
        for k in ("description", "interruptionIndex", "emotionalBalance", "empathyIndex",
                  "speechSpeedVariability", "questionsToAnswersRatio"):
            if s.get(k) not in (None, ""):
                val = _level_val(s[k], lang) if k == "speechSpeedVariability" else s[k]
                kv.append((L(k), val))
        dd = s.get("dominanceDistribution")
        if isinstance(dd, dict) and dd:
            kv.append((L("dominanceDistribution"),
                       ", ".join(f"{k}: {v}%" for k, v in dd.items())))
        if kv:
            blocks.append(("kv", kv))

    c = a.get("category")
    if isinstance(c, dict) and c:
        blocks.append(("h", 2, L("category")))
        kv = []
        if c.get("category"):
            kv.append((L("category"), c["category"]))
        if c.get("tags"):
            kv.append((L("tags"), ", ".join(map(str, c["tags"]))))
        if c.get("description"):
            kv.append((L("description"), c["description"]))
        if kv:
            blocks.append(("kv", kv))

    add_items("risks", "risks", ("severity", "impact", "status"), L, "description")
    # quotes need a quoted paragraph rather than a heading
    quotes = a.get("quotes") or []
    if quotes:
        blocks.append(("h", 2, L("quotes")))
        for it in quotes:
            if not isinstance(it, dict):
                blocks.append(("p", str(it))); continue
            blocks.append(("p", f"«{it.get('text', '')}»"))
            kv = [(L(k), it[k]) for k in ("speaker", "context") if it.get(k)]
            if kv:
                blocks.append(("kv", kv))
    add_items("technologies", "technologies", ("category", "context"), L, "name", numbered=False)
    add_items("questions", "questions", ("category", "priority", "owner"), L, "question")
    add_items("recommendations", "recommendations", ("category", "priority", "impact"), L,
              "recommendation")
    add_items("followupQuestions", "followupQuestions", ("category", "priority", "context"), L,
              "question")

    fp = a.get("formalProtocol")
    if isinstance(fp, dict) and fp:
        blocks.append(("h", 2, L("formalProtocol")))
        kv = [(L(k), fp[k]) for k in ("protocolNumber", "date", "time", "location",
              "chairman", "secretary", "nextMeeting") if fp.get(k)]
        if kv:
            blocks.append(("kv", kv))
        if fp.get("participants"):
            blocks.append(("h", 3, L("participants")))
            blocks.append(("ul", [str(x) for x in fp["participants"]]))
        if fp.get("agenda"):
            blocks.append(("h", 3, L("agenda")))
            blocks.append(("ol", [str(x) for x in fp["agenda"]]))
        if fp.get("decisions"):
            blocks.append(("h", 3, L("decisions")))
            for d in fp["decisions"]:
                if not isinstance(d, dict):
                    blocks.append(("p", str(d))); continue
                blocks.append(("h", 4, f"{L('decision')} {d.get('number', '')}".strip()))
                if d.get("text"):
                    blocks.append(("p", d["text"]))
                if d.get("votingResult"):
                    blocks.append(("kv", [(L("votingResult"), d["votingResult"])]))
        if fp.get("actionItems"):
            blocks.append(("h", 3, L("actionItems")))
            for i, it in enumerate(fp["actionItems"], 1):
                if not isinstance(it, dict):
                    blocks.append(("p", str(it))); continue
                blocks.append(("h", 4, f"{i}. {it.get('task', '')}"))
                kv = [(L(k), it[k]) for k in ("assignee", "deadline") if it.get(k)]
                if kv:
                    blocks.append(("kv", kv))
        if fp.get("protocolText"):
            blocks.append(("h", 3, L("protocolText")))
            blocks.append(("p", fp["protocolText"]))

    return blocks


def build_blocks(kind: str, data, meta: dict = None) -> list:
    meta = meta or {}
    lang = meta.get("language", "ru")
    lang = lang if lang in LBL else "ru"
    blocks = [("h", 1, _TITLES[lang][kind])]
    head = _HEAD[lang]
    head_kv = []
    if meta.get("video_name"):
        head_kv.append((head["video"], meta["video_name"]))
    head_kv.append((head["date"], meta.get("export_date") or datetime.now().strftime("%Y-%m-%d %H:%M")))
    if meta.get("duration"):
        head_kv.append((head["duration"], meta["duration"]))
    head_kv.append((head["version"], str(meta.get("version", 1))))
    blocks.append(("kv", head_kv))

    if kind == "summary":
        blocks += markdown_to_blocks(data if isinstance(data, str) else "")
    elif kind == "raw":
        blocks += build_raw_blocks(data if isinstance(data, str) else "")
    else:
        blocks += build_analysis_blocks(data if isinstance(data, dict) else {}, lang, meta)

    blocks.append(("rule",))
    blocks.append(("footer", footer_text()))
    return blocks


# -- helpers ---------------------------------------------------------------
def _esc(s) -> str:
    return _html.escape(str(s), quote=False)


def _summary_sections(text: str) -> list:
    sections, cur = [], None
    for line in (text or "").split("\n"):
        s = line.strip()
        if not s:
            if cur is not None:
                cur["content"].append("")
            continue
        if s.startswith("#"):
            if cur:
                sections.append(cur)
            n = len(s) - len(s.lstrip("#"))
            cur = {"title": s[n:].strip(), "level": min(n, 3), "content": []}
        else:
            if cur is None:
                cur = {"title": "", "level": 0, "content": []}
            cur["content"].append(s)
    if cur:
        sections.append(cur)
    for sec in sections:
        sec["content"] = "\n".join(sec["content"]).strip()
    return sections


# -- writers (consume the shared block model) ------------------------------
def _write_txt(blocks, path, meta) -> None:
    out = []
    for b in blocks:
        kind = b[0]
        if kind == "h":
            text = b[2]
            if b[1] == 1:
                out += ["", text.upper(), "=" * max(3, len(text)), ""]
            elif b[1] == 2:
                out += ["", text, "-" * max(3, len(text))]
            else:
                out += ["", text]
        elif kind == "p":
            out += [b[1], ""]
        elif kind == "ul":
            out += [f"  • {i}" for i in b[1]] + [""]
        elif kind == "ol":
            out += [f"  {n}. {i}" for n, i in enumerate(b[1], 1)] + [""]
        elif kind == "kv":
            out += [f"{k}: {v}" for k, v in b[1]] + [""]
        elif kind == "rule":
            out += ["", "-" * 48]
        elif kind == "footer":
            out += [b[1]]
    Path(path).write_text("\n".join(out).strip() + "\n", encoding="utf-8")


def _write_md(blocks, path, meta) -> None:
    out = []
    for b in blocks:
        kind = b[0]
        if kind == "h":
            out.append("#" * min(b[1], 6) + " " + b[2])
        elif kind == "p":
            out.append(b[1])
        elif kind == "ul":
            out.append("\n".join(f"- {i}" for i in b[1]))
        elif kind == "ol":
            out.append("\n".join(f"{n}. {i}" for n, i in enumerate(b[1], 1)))
        elif kind == "kv":
            out.append("\n".join(f"**{k}:** {v}" for k, v in b[1]))
        elif kind == "rule":
            out.append("---")
        elif kind == "footer":
            out.append(f"*{b[1]}*")
    Path(path).write_text("\n\n".join(out).strip() + "\n", encoding="utf-8")


_HTML_CSS = """
body{font-family:-apple-system,'Segoe UI',Roboto,Arial,sans-serif;line-height:1.6;
color:#1f2329;background:#f5f6f8;margin:0;padding:24px}
.container{max-width:900px;margin:0 auto;background:#fff;padding:40px;
border-radius:8px;box-shadow:0 2px 10px rgba(0,0,0,.08)}
h1{font-size:2em;color:#1a2533;margin:0 0 24px}
h2{font-size:1.4em;color:#22303f;margin:28px 0 12px;padding-bottom:8px;border-bottom:2px solid #3498db}
h3{font-size:1.12em;color:#2c3e50;margin:18px 0 8px}
h4{font-size:1em;color:#2c3e50;margin:14px 0 6px}
p{margin:0 0 12px}
ul,ol{margin:0 0 14px 28px}
li{margin:4px 0}
.kv{display:flex;gap:10px;margin:4px 0;font-size:.95em}
.kv .k{font-weight:600;color:#34495e;min-width:160px}
.kv .v{color:#4a5568}
hr{border:none;border-top:1px solid #e2e6ea;margin:32px 0 16px}
.footer{text-align:center;color:#95a5a6;font-size:.9em}
"""


def _write_html(blocks, path, meta) -> None:
    parts = []
    for b in blocks:
        kind = b[0]
        if kind == "h":
            parts.append(f"<h{b[1]}>{_esc(b[2])}</h{b[1]}>")
        elif kind == "p":
            parts.append(f"<p>{_esc(b[1]).replace(chr(10), '<br>')}</p>")
        elif kind == "ul":
            parts.append("<ul>" + "".join(f"<li>{_esc(i)}</li>" for i in b[1]) + "</ul>")
        elif kind == "ol":
            parts.append("<ol>" + "".join(f"<li>{_esc(i)}</li>" for i in b[1]) + "</ol>")
        elif kind == "kv":
            rows = "".join(
                f'<div class="kv"><span class="k">{_esc(k)}</span>'
                f'<span class="v">{_esc(v)}</span></div>' for k, v in b[1])
            parts.append(rows)
        elif kind == "rule":
            parts.append("<hr>")
        elif kind == "footer":
            parts.append(f'<div class="footer">{_esc(b[1])}</div>')
    title = blocks[0][2] if blocks and blocks[0][0] == "h" else APP_NAME
    doc = (f'<!DOCTYPE html>\n<html lang="{meta.get("language", "ru")}">\n<head>\n'
           f'<meta charset="UTF-8">\n<meta name="viewport" content="width=device-width,'
           f' initial-scale=1.0">\n<title>{_esc(title)}</title>\n<style>{_HTML_CSS}</style>\n'
           f'</head>\n<body>\n<div class="container">\n' + "\n".join(parts)
           + "\n</div>\n</body>\n</html>\n")
    Path(path).write_text(doc, encoding="utf-8")


def _register_pdf_font():
    import os
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    # Prefer a broad Unicode face on Windows.  Arial/Segoe render Cyrillic but
    # silently drop CJK glyphs that may legitimately occur in names, quotes or
    # model output, making PDF content differ from TXT/HTML/DOCX/JSON.
    candidates = ["C:/Windows/Fonts/malgun.ttf",
                  "C:/Windows/Fonts/arial.ttf", "C:/Windows/Fonts/segoeui.ttf",
                  "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
                  "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"]
    bold = {"C:/Windows/Fonts/malgun.ttf": "C:/Windows/Fonts/malgunbd.ttf",
            "C:/Windows/Fonts/arial.ttf": "C:/Windows/Fonts/arialbd.ttf",
            "C:/Windows/Fonts/segoeui.ttf": "C:/Windows/Fonts/segoeuib.ttf"}
    import time
    for path in candidates:
        if not os.path.exists(path):
            continue
        # A font file can be momentarily unreadable (another process has it open);
        # one retry is the difference between a correct PDF and a silent fallback.
        for attempt in range(2):
            try:
                pdfmetrics.registerFont(TTFont("Body", path))
                bp = bold.get(path)
                if bp and os.path.exists(bp):
                    pdfmetrics.registerFont(TTFont("Body-Bold", bp))
                    pdfmetrics.registerFontFamily("Body", normal="Body", bold="Body-Bold")
                return "Body"
            except Exception:  # noqa: BLE001 - try the next candidate/attempt
                if attempt == 0:
                    time.sleep(0.2)
    # Helvetica has NO Cyrillic glyphs. Returning it silently is how a PDF ends up
    # missing every Russian word while txt/md/html/docx are fine - the caller must
    # decide, so it is never mistaken for a successfully exported document.
    return "Helvetica"


def _rl(text) -> str:
    return _html.escape(str(text), quote=False).replace("\n", "<br/>")


def _block_text(block) -> str:
    """Every string a block carries, for checks that must see all of its text."""
    out = []
    for part in block[1:]:
        if isinstance(part, str):
            out.append(part)
        elif isinstance(part, (list, tuple)):
            for item in part:
                if isinstance(item, str):
                    out.append(item)
                elif isinstance(item, (list, tuple)):
                    out.extend(str(x) for x in item)
    return " ".join(out)


def _write_pdf(blocks, path, meta) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    font = _register_pdf_font()
    if font == "Helvetica" and any(
            ord(ch) > 127 for b in blocks for ch in _block_text(b)):
        raise RuntimeError(
            "PDF export needs a Unicode font, and none could be registered "
            "(tried Malgun/Arial/Segoe UI/Noto/DejaVu). Helvetica cannot render "
            "Cyrillic, so the PDF would silently lose the text. Install one of "
            "those fonts, or export to DOCX/HTML/MD instead.")
    base = getSampleStyleSheet()
    sizes = {1: 22, 2: 16, 3: 13, 4: 11.5}
    styles = {lvl: ParagraphStyle(f"H{lvl}", parent=base["Heading1"], fontName=font,
                                  fontSize=sz, spaceBefore=10, spaceAfter=6,
                                  textColor="#22303f") for lvl, sz in sizes.items()}
    body = ParagraphStyle("Body", parent=base["Normal"], fontName=font, fontSize=11,
                          leading=15, spaceAfter=8)
    kv_style = ParagraphStyle("Kv", parent=body, leftIndent=10, spaceAfter=2)
    footer_style = ParagraphStyle("Foot", parent=body, fontSize=9, textColor="#95a5a6",
                                  alignment=1, spaceBefore=8)
    story = []
    for b in blocks:
        kind = b[0]
        if kind == "h":
            story.append(Paragraph(_rl(b[2]), styles.get(b[1], styles[3])))
        elif kind == "p":
            story.append(Paragraph(_rl(b[1]), body))
        elif kind == "ul":
            for i in b[1]:
                story.append(Paragraph("• " + _rl(i), kv_style))
            story.append(Spacer(1, 0.08 * inch))
        elif kind == "ol":
            for n, i in enumerate(b[1], 1):
                story.append(Paragraph(f"{n}. " + _rl(i), kv_style))
            story.append(Spacer(1, 0.08 * inch))
        elif kind == "kv":
            for k, v in b[1]:
                story.append(Paragraph(f"<b>{_rl(k)}:</b> {_rl(v)}", kv_style))
            story.append(Spacer(1, 0.06 * inch))
        elif kind == "rule":
            story.append(Spacer(1, 0.2 * inch))
        elif kind == "footer":
            story.append(Paragraph(_rl(b[1]), footer_style))
    SimpleDocTemplate(str(path), pagesize=A4, rightMargin=54, leftMargin=54,
                      topMargin=54, bottomMargin=40).build(story)


def _write_docx(blocks, path, meta) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    for b in blocks:
        kind = b[0]
        if kind == "h":
            doc.add_heading(b[2], level=min(b[1], 4))
        elif kind == "p":
            doc.add_paragraph(b[1])
        elif kind == "ul":
            for i in b[1]:
                doc.add_paragraph(str(i), style="List Bullet")
        elif kind == "ol":
            for i in b[1]:
                doc.add_paragraph(str(i), style="List Number")
        elif kind == "kv":
            for k, v in b[1]:
                p = doc.add_paragraph()
                run = p.add_run(f"{k}: ")
                run.bold = True
                p.add_run(str(v))
        elif kind == "rule":
            doc.add_paragraph("—" * 24)
        elif kind == "footer":
            p = doc.add_paragraph()
            run = p.add_run(b[1])
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)
    doc.save(str(path))


def _write_json(kind, data, path, meta) -> None:
    metadata = {
        "video_name": meta.get("video_name", ""),
        "export_date": meta.get("export_date") or datetime.now().isoformat(),
        "duration": meta.get("duration", ""),
        "version": meta.get("version", 1),
        "app_version": meta.get("app_version", app_version()),
        "kind": kind,
    }
    payload = {"_generator": footer_text(), "metadata": metadata}
    if kind == "analysis":
        payload["analysis"] = data if isinstance(data, dict) else {}
    elif kind == "summary":
        text = data if isinstance(data, str) else ""
        payload["summary"] = {"full_text": text, "sections": _summary_sections(text)}
    else:
        payload["transcript"] = data if isinstance(data, str) else ""
    Path(path).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


_ANALYSIS_CSS = """
*{margin:0;padding:0;box-sizing:border-box}
body{font-family:-apple-system,'Segoe UI',Roboto,sans-serif;background:#f5f6f8;color:#333;
padding:2rem;line-height:1.6}
.container{max-width:1100px;margin:0 auto}
h1{font-size:2.2rem;margin-bottom:.5rem;color:#1a1a1a;border-bottom:3px solid #007acc;padding-bottom:.5rem}
.meta{color:#666;margin-bottom:2rem;font-size:.95rem;background:#fff;padding:1rem;border-radius:8px;border-left:4px solid #007acc}
.section{background:#fff;border:1px solid #e2e6ea;border-radius:12px;padding:1.75rem;margin-bottom:1.5rem;box-shadow:0 2px 8px rgba(0,0,0,.06)}
.section h2{font-size:1.4rem;margin-bottom:1.25rem;color:#007acc;border-bottom:2px solid #eef1f4;padding-bottom:.5rem}
.grid{display:grid;grid-template-columns:repeat(3,minmax(0,1fr));gap:1.1rem;margin-bottom:1rem;align-items:start}
.card{background:linear-gradient(135deg,#f8f9fa 0%,#eef1f4 100%);border:1px solid #dee2e6;border-radius:10px;padding:1.1rem;display:flex;align-items:flex-start;gap:.9rem;min-width:0}
.card.full{grid-column:1/-1;background:#f8f9fa}
.icon{font-size:1.8rem}.clabel{color:#666;font-size:.8rem;text-transform:uppercase;letter-spacing:.5px;font-weight:600}
.cvalue{font-size:1.15rem;font-weight:700;color:#1a1a1a;overflow-wrap:anywhere;word-break:normal}
.tags{display:flex;flex-wrap:wrap;gap:.5rem;margin-top:.5rem}
.tag{padding:.35rem .85rem;background:linear-gradient(135deg,#007acc,#4a9eff);color:#fff;border-radius:16px;font-size:.85rem;font-weight:600}
.tag.soft{background:#e8f5e9;color:#2e7d32;border:1px solid #a5d6a7}
.item{background:#f8f9fa;border:1px solid #dee2e6;border-left:4px solid #007acc;border-radius:8px;padding:1.1rem;margin-bottom:.9rem}
.item.sev-high{border-left-color:#dc3545}.item.sev-medium{border-left-color:#ffc107}.item.sev-low{border-left-color:#28a745}
.ihead{display:flex;justify-content:space-between;gap:1rem;align-items:flex-start;margin-bottom:.6rem}
.itask{font-weight:600;font-size:1.03rem;color:#1a1a1a;flex:1}
.badge{padding:.3rem .8rem;border-radius:14px;font-size:.75rem;font-weight:700;text-transform:uppercase;white-space:nowrap}
.b-high{background:#dc3545;color:#fff}.b-medium{background:#ffc107;color:#000}.b-low{background:#28a745;color:#fff}
.imeta{display:flex;gap:1.3rem;font-size:.88rem;color:#666;flex-wrap:wrap}
.bar{height:9px;background:#e6e9ec;border-radius:5px;overflow:hidden;margin-top:.5rem}
.fill{height:100%;background:linear-gradient(90deg,#007acc,#4a9eff)}
.srow{background:#f8f9fa;border:1px solid #dee2e6;border-radius:8px;padding:1rem;margin-bottom:.7rem}
.srow .clabel{margin-bottom:.3rem}.sval{font-weight:700;font-size:1.15rem}
.pos{color:#28a745}.neu{color:#6c757d}.neg{color:#dc3545}
.badge-lg{display:inline-block;padding:.7rem 1.4rem;background:linear-gradient(135deg,#007acc,#4a9eff);color:#fff;border-radius:8px;font-weight:700;font-size:1.05rem;margin-bottom:1rem}
blockquote{border-left:4px solid #007acc;background:#f8f9fa;margin:0 0 .8rem;padding:.7rem 1rem;border-radius:0 8px 8px 0;font-style:italic}
blockquote .who{display:block;font-style:normal;color:#007acc;font-weight:600;margin-top:.4rem;font-size:.85rem}
.desc{margin-top:.8rem;font-size:.93rem;color:#444}
ul{margin:.4rem 0 .4rem 1.4rem}li{margin:.3rem 0}
.footer{text-align:center;margin-top:2rem;color:#95a5a6;font-size:.9rem}
@media(max-width:760px){.grid{grid-template-columns:1fr}.card.full{grid-column:auto}}
"""

_PRIO = {"ru": {"high": "ВЫСОКИЙ", "medium": "СРЕДНИЙ", "low": "НИЗКИЙ"},
         "en": {"high": "HIGH", "medium": "MEDIUM", "low": "LOW"}}


def _write_analysis_html(a: dict, path, meta: dict) -> None:
    """Rich, card-based analysis report (ports the original app's HTML look:
    characteristic cards, priority badges, sentiment bars, category tags)."""
    lang = meta.get("language", "ru")
    L = lambda k: LBL.get(lang, LBL["en"]).get(k, k)   # noqa: E731
    prio = _PRIO.get(lang, _PRIO["en"])
    a = a or {}
    S: list = []

    def sect(title):
        S.append(f'<div class="section"><h2>{title}</h2>')

    def end():
        S.append("</div>")

    def prio_badge(p):
        p = str(p or "").lower()
        cls = {"high": "b-high", "medium": "b-medium", "low": "b-low"}.get(p, "b-medium")
        return f'<span class="badge {cls}">{_esc(prio.get(p, p or "—"))}</span>' if p else ""

    def item_cards(items, task_key, sev_key=None, meta_keys=()):
        for it in items:
            if not isinstance(it, dict):
                S.append(f'<div class="item"><div class="itask">{_esc(it)}</div></div>'); continue
            sev = str(it.get(sev_key, "") if sev_key else it.get("priority", "")).lower()
            sev_cls = f" sev-{sev}" if sev in ("high", "medium", "low") else ""
            metas = "".join(f"<span>{_esc(L(k))}: {_esc(it.get(k))}</span>"
                            for k in meta_keys if it.get(k) not in (None, ""))
            S.append(
                f'<div class="item{sev_cls}"><div class="ihead">'
                f'<div class="itask">{_esc(it.get(task_key, ""))}</div>{prio_badge(sev)}</div>'
                f'<div class="imeta">{metas}</div></div>')

    # -- Характеристики -------------------------------------------------
    ch = a.get("characteristics") or {}
    kt = ch.get("keyTopics") or []
    fw = ch.get("frequentWords") or []
    dur = meta.get("duration")
    words = meta.get("wordCount")
    parts = meta.get("participants")
    sect("📋 " + L("characteristics"))
    S.append('<div class="grid">')
    characteristic_cards = []
    if dur not in (None, ""):
        characteristic_cards.append(("⏱️", L("duration"), dur))
    if parts not in (None, ""):
        characteristic_cards.append(("👥", L("participants"), parts))
    if words is not None:
        word_value = f"{int(words):,}".replace(",", " ")
        characteristic_cards.append(("📝", L("wordCount"), word_value))
    for icon, lab, val in characteristic_cards:
        S.append(f'<div class="card"><div class="icon">{icon}</div><div>'
                 f'<div class="clabel">{_esc(lab)}</div><div class="cvalue">{_esc(val)}</div></div></div>')
    if kt:
        tags = "".join(f'<span class="tag">{_esc(t)}</span>' for t in kt)
        S.append(f'<div class="card full"><div class="icon">🔑</div><div style="flex:1">'
                 f'<div class="clabel">{_esc(L("keyTopics"))}</div><div class="tags">{tags}</div></div></div>')
    if fw:
        values = [
            f"{x[0]} — {x[1]}" if isinstance(x, (list, tuple)) and len(x) >= 2 else str(x)
            for x in fw
        ]
        items = "".join(f"<li>{_esc(value)}</li>" for value in values)
        S.append(f'<div class="card full"><div class="icon">🔤</div><div style="flex:1">'
                 f'<div class="clabel">{_esc(L("frequentWords"))}</div><ul>{items}</ul></div></div>')
    S.append("</div>")
    end()

    # -- Action items ---------------------------------------------------
    ai = a.get("actionItems") or []
    if ai:
        sect("✅ " + L("actionItems"))
        item_cards(ai, "task", "priority", ("assignee", "deadline"))
        end()

    # -- Sentiment ------------------------------------------------------
    s = a.get("sentiment")
    if isinstance(s, dict) and s:
        sect("😊 " + L("sentiment"))
        ov = str(s.get("overall", "")).lower()
        ovcls = {"positive": "pos", "neutral": "neu", "negative": "neg"}.get(ov, "neu")
        S.append(f'<div class="srow"><div class="clabel">{_esc(L("overall"))}</div>'
                 f'<div class="sval {ovcls}">{_esc(_sentiment_val(s.get("overall", "—"), lang))}</div></div>')
        if s.get("engagement") not in (None, ""):
            S.append(f'<div class="srow"><div class="clabel">{_esc(L("engagement"))}</div>'
                     f'<div class="sval">{_esc(_level_val(s.get("engagement"), lang))}</div></div>')
        if "hasConflict" in s:
            conflict = L("yes") if s.get("hasConflict") else L("no")
            S.append(f'<div class="srow"><div class="clabel">{_esc(L("hasConflict"))}</div>'
                     f'<div class="sval">{_esc(conflict)}</div></div>')
        for k in ("interruptionIndex", "emotionalBalance", "empathyIndex"):
            v = s.get(k)
            if isinstance(v, (int, float)):
                w = max(0, min(100, int(v)))
                S.append(f'<div class="srow"><div class="clabel">{_esc(L(k))}</div>'
                         f'<div class="sval">{w}/100</div>'
                         f'<div class="bar"><div class="fill" style="width:{w}%"></div></div></div>')
        for k in ("speechSpeedVariability", "questionsToAnswersRatio"):
            if s.get(k) not in (None, ""):
                val = _level_val(s.get(k), lang) if k == "speechSpeedVariability" else s.get(k)
                S.append(f'<div class="srow"><div class="clabel">{_esc(L(k))}</div>'
                         f'<div class="sval">{_esc(val)}</div></div>')
        if s.get("emotions"):
            tags = "".join(f'<span class="tag soft">{_esc(e)}</span>' for e in s["emotions"])
            S.append(f'<div class="srow"><div class="clabel">{_esc(L("emotions"))}</div>'
                     f'<div class="tags">{tags}</div></div>')
        dd = s.get("dominanceDistribution")
        if isinstance(dd, dict) and dd:
            S.append(f'<div class="srow"><div class="clabel">{_esc(L("dominanceDistribution"))}</div>')
            for who, pc in dd.items():
                try:
                    w = max(0, min(100, int(float(pc))))
                except (TypeError, ValueError):
                    w = 0
                S.append(f'<div style="margin-top:.5rem"><div style="display:flex;'
                         f'justify-content:space-between"><span>{_esc(who)}</span>'
                         f'<b>{_esc(pc)}%</b></div><div class="bar"><div class="fill" '
                         f'style="width:{w}%"></div></div></div>')
            S.append("</div>")
        if s.get("description"):
            S.append(f'<div class="desc">{_esc(s["description"])}</div>')
        end()

    # -- Category -------------------------------------------------------
    c = a.get("category")
    if isinstance(c, dict) and c:
        sect("📂 " + L("category"))
        if c.get("category"):
            S.append(f'<div class="badge-lg">{_esc(c["category"])}</div>')
        if c.get("tags"):
            tags = "".join(f'<span class="tag soft">{_esc(t)}</span>' for t in c["tags"])
            S.append(f'<div class="tags">{tags}</div>')
        if c.get("description"):
            S.append(f'<div class="desc">{_esc(c["description"])}</div>')
        end()

    # -- Risks / questions / recommendations / follow-ups (card lists) --
    for key, emoji, task, mk in (
            ("risks", "🔴", "description", ("severity", "impact", "status")),
            ("questions", "❓", "question", ("category", "priority", "owner")),
            ("recommendations", "💡", "recommendation", ("category", "priority", "impact")),
            ("followupQuestions", "🔄", "question", ("category", "priority", "context"))):
        items = a.get(key) or []
        if items:
            sect(f"{emoji} " + L(key))
            item_cards(items, task, "severity" if key == "risks" else "priority", mk)
            end()

    # -- Quotes ---------------------------------------------------------
    quotes = a.get("quotes") or []
    if quotes:
        sect("💬 " + L("quotes"))
        for q in quotes:
            if isinstance(q, dict):
                who = " · ".join(str(q.get(k)) for k in ("speaker", "context") if q.get(k))
                S.append(f'<blockquote>«{_esc(q.get("text", ""))}»'
                         + (f'<span class="who">{_esc(who)}</span>' if who else "") + "</blockquote>")
            else:
                S.append(f'<blockquote>«{_esc(q)}»</blockquote>')
        end()

    # -- Technologies ---------------------------------------------------
    tech = a.get("technologies") or []
    if tech:
        sect("💻 " + L("technologies"))
        for t in tech:
            if isinstance(t, dict):
                metas = " · ".join(_esc(t.get(k)) for k in ("category", "context") if t.get(k))
                S.append(f'<div class="item"><div class="itask">{_esc(t.get("name", ""))}</div>'
                         f'<div class="imeta"><span>{metas}</span></div></div>')
            else:
                S.append(f'<div class="item"><div class="itask">{_esc(t)}</div></div>')
        end()

    # -- Formal protocol -------------------------------------------------
    # Rendered in FULL (metadata, participants, agenda, decisions, action items,
    # protocol text): this is an official document — dropping any part of it
    # would silently lose data the user must keep.
    fp = a.get("formalProtocol")
    if isinstance(fp, dict) and fp:
        sect("📜 " + L("formalProtocol"))
        kv = [(L(k), fp[k]) for k in ("protocolNumber", "date", "time", "location",
                                      "chairman", "secretary", "nextMeeting") if fp.get(k)]
        for label, value in kv:
            S.append(f'<div class="srow"><div class="clabel">{_esc(label)}</div>'
                     f'<div class="sval">{_esc(value)}</div></div>')
        if fp.get("participants"):
            tags = "".join(f'<span class="tag soft">{_esc(p)}</span>' for p in fp["participants"])
            S.append(f'<div class="srow"><div class="clabel">{_esc(L("participants"))}</div>'
                     f'<div class="tags">{tags}</div></div>')
        if fp.get("agenda"):
            items = "".join(f"<li>{_esc(x)}</li>" for x in fp["agenda"])
            S.append(f'<div class="clabel">{_esc(L("agenda"))}</div><ol>{items}</ol>')
        if fp.get("decisions"):
            S.append(f'<div class="clabel">{_esc(L("decisions"))}</div>')
            for d in fp["decisions"]:
                if not isinstance(d, dict):
                    S.append(f'<div class="item"><div class="itask">{_esc(d)}</div></div>')
                    continue
                head = f"{L('decision')} {d.get('number', '')}".strip()
                vote = (f'<div class="imeta"><span>{_esc(L("votingResult"))}: '
                        f'{_esc(d.get("votingResult"))}</span></div>') if d.get("votingResult") else ""
                S.append(f'<div class="item"><div class="ihead"><div class="itask">'
                         f'{_esc(head)}: {_esc(d.get("text", ""))}</div></div>{vote}</div>')
        if fp.get("actionItems"):
            S.append(f'<div class="clabel">{_esc(L("actionItems"))}</div>')
            item_cards(fp["actionItems"], "task", None, ("assignee", "deadline"))
        if fp.get("protocolText"):
            S.append(f'<div class="clabel">{_esc(L("protocolText"))}</div>'
                     f'<div class="desc" style="white-space:pre-wrap">'
                     f'{_esc(fp["protocolText"])}</div>')
        end()

    title = L("analysisReport")
    body = "\n".join(S)
    head = _HEAD.get(lang, _HEAD["en"])
    head_rows = []
    if meta.get("video_name"):
        head_rows.append((head["video"], meta["video_name"]))
    head_rows.append((head["date"], meta.get("export_date")
                      or datetime.now().strftime("%Y-%m-%d %H:%M")))
    if meta.get("duration"):
        head_rows.append((head["duration"], meta["duration"]))
    head_rows.append((head["version"], str(meta.get("version", 1))))
    meta_html = "".join(
        f'<div><b>{_esc(label)}:</b> {_esc(value)}</div>' for label, value in head_rows
    )
    doc = (f'<!DOCTYPE html>\n<html lang="{lang}"><head><meta charset="UTF-8">'
           f'<meta name="viewport" content="width=device-width,initial-scale=1.0">'
           f'<title>{_esc(title)} — {_esc(meta.get("video_name", ""))}</title>'
           f'<style>{_ANALYSIS_CSS}</style></head><body><div class="container">'
           f'<h1>📊 {_esc(title)}</h1>'
           f'<div class="meta">{meta_html}</div>'
           f'{body}<div class="footer">{_esc(footer_text())}</div>'
           f'</div></body></html>\n')
    Path(path).write_text(doc, encoding="utf-8")


_WRITERS = {"txt": _write_txt, "md": _write_md, "html": _write_html,
            "pdf": _write_pdf, "docx": _write_docx}


def export(kind: str, data, fmt: str, out_path, meta: dict = None) -> str:
    """Export one artifact to one format. Returns the written path."""
    if kind not in KINDS:
        raise ValueError(f"Unknown kind: {kind} (expected {KINDS})")
    if fmt not in FORMATS:
        raise ValueError(f"Unknown format: {fmt} (expected {FORMATS})")
    meta = dict(meta or {})
    meta.setdefault("app_version", app_version())
    out_path = str(out_path)
    if fmt == "json":
        _write_json(kind, data, out_path, meta)
    elif kind == "analysis" and fmt == "html":
        # Rich, card-based analysis report (matches the original app's look).
        _write_analysis_html(data if isinstance(data, dict) else {}, out_path, meta)
    else:
        _WRITERS[fmt](build_blocks(kind, data, meta), out_path, meta)
    return out_path
